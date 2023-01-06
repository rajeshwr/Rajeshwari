#!/usr/bin/env python
# coding: utf-8

# 1. In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened?

# For PdfFileReader() file objects should be opened in rb -> read binary mode, Whereas for PdfFileWriter() file objects should be opened in wb -> write binary mode.

# 2. From a PdfFileReader object, how do you get a Page object for page 5?

# PdfFileReader class provides a method called getPage(page_no) to get a page object. Calling getPage(4) will return a Page object for page 5 since page 0 is the first page

# In[ ]:


import PyPDF2 as pdf
pdfFileObj = open("cardealer.pdf",'rb')
pdfReader = pdf.PdfFileReader(pdfFileObj)
pageObj = pdfReader.getPage(4)
pageObj.extractText()    


# 3. What PdfFileReader variable stores the number of pages in the PDF document?

# In[ ]:


#getNumPages() method of PdfFileReader class stores the no pages in a PDF document
import PyPDF2 as pdf
pdfFileObj = open("cardealer.pdf",'rb')
pdfReader = pdf.PdfFileReader(pdfFileObj)
print(pdfReader.numPages)


# 4. If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do before you can obtain Page objects from it?

# In[ ]:


#Before we obtain the page object, the pdf has to be decrypted by calling .decrypt('swordfish')
from PyPDF2 import PdfFileReader
pdf_reader = PdfFileReader(file_path)
if pdf_reader.isEncrypted: # to check whether the pdf is encrypted or not
    pdf_reader.decrypt("swordfish")
for page in pdf_reader.pages:
    print(page.extractText()) # to print the text data of a page from pdf


# 5. What methods do you use to rotate a page?

# PyPDF2 Package provides 2 methods to rotate a page:
# 
# rotateClockWise() -> For Clockwise rotation
# rotateCounterClockWise() -> For Counter Clockwise rotation
# The PyPDF2 package only allows you to rotate a page in increments of 90 degrees. You will receive an AssertionError otherwise.

# 6. What is the difference between a Run object and a Paragraph object?

# Paragraph Object : A document contains multiple paragraphs. A paragraph begins on a new line and contains multiple
# runs. The Document object contains a list of Paragraph objects for the paragraphs in the document. (A new paragraph begins whenever the user presses ENTER or RETURN while typing in a Word document.)
# 
# Run Objects : Runs are contiguous groups of characters within a paragraph with the same style

# 7. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc?

# In[ ]:


from docx import Document
doc = Document("sample_file.docx") # Path of the Docx file
print(doc.paragraphs) # Prints the list of Paragraph objects for a Document
for paragraph in doc.paragraphs:
    print(paragraph.text) # Prints the text in the paragraph


# 8. What type of object has bold, underline, italic, strike, and outline variables?

# A Run object has bold, underline,italic,strike and outline variables

# 9. What is the difference between False, True, and None for the bold variable?

# Runs can be further styled using text attributes. Each attribute can be set to one of three values:
# 
# True (the attribute is always enabled, no matter what other styles are applied to the run),
# False (the attribute is always disabled),
# None (defaults to whatever the run’s style is set to)
# 
# True always makes the Run object bolded and False makes it always not bolded, no matter what the style’s bold setting is. None will make the Run object just use the style’s bold setting

# 10. How do you create a Document object for a new Word document?

# In[ ]:


#By Calling the docx.Document() function
from docx import Document
document = Document()
document.add_paragraph("Learn python programming")
document.save('mydocument.docx')


# 11. How do you add a paragraph with the text 'Hello, there!' to a Document object stored in a variable named doc?

# In[ ]:


import docx
doc = docx.Document()
doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')


# 12. What integers represent the levels of headings available in Word documents?

# integer from 0 to 4
# The arguments to add_heading() are a string of the heading text and an integer from 0 to 4. The integer 0 makes the heading the Title style, which is used for the top of the document. Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the lowest subheading
